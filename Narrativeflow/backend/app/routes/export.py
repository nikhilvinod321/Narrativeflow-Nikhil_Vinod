"""
Export Routes - Export stories in various formats
"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import Optional
from uuid import UUID
import io
import os
import mimetypes
from urllib.parse import quote, urlparse
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ebooklib import epub
import html
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.utils import ImageReader

from app.database import get_db
from app.models.plotline import Plotline
from app.models.character import Character
from app.services.story_service import StoryService
from app.services.chapter_service import ChapterService

router = APIRouter()

story_service = StoryService()
chapter_service = ChapterService()

STATIC_ROOT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "static"
)



def clean_and_parse_html(content: str):
    """Clean encoded HTML content and return soup object"""
    if not content:
        return None, False
    
    raw = content
    # Handle JSON-style escaped angle brackets if present
    if '\\u003c' in raw or '\\u003e' in raw:
        raw = raw.replace('\\u003c', '<').replace('\\u003e', '>')
    # Unescape up to 3 times to handle double-encoding
    for _ in range(3):
        # We check broadly for entities or just unescape blindly to be safe
        # checking "if '&lt;' in raw" might fail if it's mixed with other entities
        # or if it's just one entity. Safest to just try unescaping.
        # But we don't want to unescape regular text unnecessarily if big.
        # A quick check for '&' is usually enough.
        if '&' in raw:
            try:
                 new_raw = html.unescape(raw)
                 if new_raw == raw:
                     break
                 raw = new_raw
            except:
                break
        else:
            break
            
    soup = BeautifulSoup(raw, 'html.parser')
    
    # Check if we have actual tags
    # If the text *starts* with < and ends with >, it's likely HTML even if soup.find fails for some reason
    has_tags = bool(soup.find())
    
    # Fallback: Check for HTML-like strings directly if soup was too strict or failed
    if not has_tags and (('<p' in raw) or ('<div' in raw) or ('<br' in raw)):
         has_tags = True

    return soup, has_tags

def extract_plain_text(content: str):
    """Return clean text from possibly-HTML content"""
    soup, _ = clean_and_parse_html(content)
    if not soup:
        return ""
    return soup.get_text(separator="\n\n")

def iter_html_children(soup):
    if not soup:
        return []
    if soup.body:
        return list(soup.body.children)
    return list(soup.children)

def resolve_image_path(src: str):
    if not src:
        return None

    # Normalize common malformed URLs
    src = src.strip().replace("local host", "localhost").replace(" ", "")
    parsed = urlparse(src)
    path = None

    if parsed.scheme and parsed.scheme.startswith("http"):
        if "localhost" in parsed.netloc or "127.0.0.1" in parsed.netloc:
            path = parsed.path
        else:
            return None
    else:
        path = parsed.path or src

    if not path:
        return None

    if path.startswith("/static/"):
        rel = path[len("/static/"):]
    elif path.startswith("static/"):
        rel = path[len("static/"):]
    else:
        rel = path.lstrip("/")

    if rel.startswith("static/"):
        rel = rel[len("static/"):]

    final_path = os.path.normpath(os.path.join(STATIC_ROOT, rel))
    return final_path if os.path.exists(final_path) else None

def build_pdf_image(image_path: str, max_width: float):
    try:
        reader = ImageReader(image_path)
        iw, ih = reader.getSize()
        if iw <= 0 or ih <= 0:
            return None
        scale = min(max_width / iw, 1.0)
        return RLImage(image_path, width=iw * scale, height=ih * scale)
    except Exception:
        return None

def build_epub_body_html(content: str, book: epub.EpubBook, image_map: dict):
    soup, is_html = clean_and_parse_html(content)
    if is_html and soup:
        body = soup.body or soup
        for img in body.find_all('img'):
            src = img.get('src')
            image_path = resolve_image_path(src) if src else None
            if not image_path:
                continue
            if image_path not in image_map:
                file_name = f"images/{os.path.basename(image_path)}"
                media_type = mimetypes.guess_type(file_name)[0] or "image/png"
                with open(image_path, "rb") as f:
                    data = f.read()
                img_item = epub.EpubItem(
                    uid=f"img_{len(image_map) + 1}",
                    file_name=file_name,
                    media_type=media_type,
                    content=data
                )
                book.add_item(img_item)
                image_map[image_path] = file_name
            img['src'] = image_map[image_path]
        return body.decode_contents()

    plain = extract_plain_text(content)
    paragraphs = [p.strip() for p in plain.split('\n') if p.strip()]
    return "".join(f"<p>{html.escape(p)}</p>" for p in paragraphs)

def html_to_markdown(html_content):
    if not html_content:
        return ""
    try:
        # Use our robust cleaner
        soup, is_html = clean_and_parse_html(html_content)
        
        if not is_html:
            # If it didn't parse as HTML tags, return cleaned text
            return extract_plain_text(html_content)

        # Simple converter
        markdown = ""
        for element in iter_html_children(soup):
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text: markdown += text + "\n\n"
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                markdown += "#" * level + " " + element.get_text().strip() + "\n\n"
            elif element.name == 'p':
                # Check for images
                img = element.find('img')
                if img and img.get('src'):
                    alt = img.get('alt') or "Image"
                    src = img.get('src')
                    markdown += f"![{alt}]({src})\n\n"
                
                text = element.get_text().strip()
                if text:
                    markdown += text + "\n\n"
            elif element.name == 'img':
                alt = element.get('alt') or "Image"
                src = element.get('src')
                if src:
                    markdown += f"![{alt}]({src})\n\n"
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    markdown += f"- {li.get_text().strip()}\n"
                markdown += "\n"
            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li')):
                    markdown += f"{i+1}. {li.get_text().strip()}\n"
                markdown += "\n"
            elif element.name == 'blockquote':
                markdown += f"> {element.get_text().strip()}\n\n"
            else:
                # Fallback
                text = element.get_text().strip()
                if text: markdown += text + "\n\n"
                
        return markdown
    except Exception as e:
        print(f"Error converting MD: {e}")
        return extract_plain_text(html_content)

@router.get("/{story_id}/markdown")
async def export_markdown(
    story_id: UUID,
    include_notes: bool = False,
    db: AsyncSession = Depends(get_db)
):
    """Export story as Markdown"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Build markdown content
    md_content = f"# {story.title}\n\n"
    
    if story.subtitle:
        md_content += f"*{story.subtitle}*\n\n"
    
    if story.logline:
        md_content += f"> {story.logline}\n\n"
    
    md_content += f"**Genre:** {story.genre.value.replace('_', ' ').title()}\n"
    md_content += f"**Tone:** {story.tone.value.title()}\n"
    md_content += f"**Word Count:** {story.word_count:,}\n\n"
    
    md_content += "---\n\n"
    
    for chapter in chapters:
        md_content += f"## Chapter {chapter.number}: {chapter.title}\n\n"
        
        if include_notes and chapter.notes:
            md_content += f"*Notes: {chapter.notes}*\n\n"
        
        if chapter.content:
            md_content += html_to_markdown(chapter.content)
        
        md_content += "---\n\n"
    
    # Return as downloadable file
    buffer = io.BytesIO(md_content.encode('utf-8'))
    
    filename = f"{story.title.replace(' ', '_')}.md"
    # Encode filename for non-ASCII characters (RFC 5987)
    encoded_filename = quote(filename, safe='')
    
    return StreamingResponse(
        buffer,
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{story_id}/text")
async def export_plain_text(
    story_id: UUID,
    db: AsyncSession = Depends(get_db)
):
    """Export story as plain text"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Build plain text content
    text_content = f"{story.title.upper()}\n"
    text_content += "=" * len(story.title) + "\n\n"
    
    if story.subtitle:
        text_content += f"{story.subtitle}\n\n"
    
    for chapter in chapters:
        text_content += f"\n\nCHAPTER {chapter.number}: {chapter.title.upper()}\n"
        text_content += "-" * 40 + "\n\n"
        
        if chapter.content:
            text_context_raw = chapter.content
            # Strip HTML if present
            try:
                cleaned = extract_plain_text(text_context_raw)
                text_content += cleaned + "\n"
            except:
                text_content += extract_plain_text(text_context_raw) + "\n"
    
    buffer = io.BytesIO(text_content.encode('utf-8'))
    filename = f"{story.title.replace(' ', '_')}.txt"
    # Encode filename for non-ASCII characters (RFC 5987)
    encoded_filename = quote(filename, safe='')
    
    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{story_id}/docx")
async def export_docx(
    story_id: UUID,
    db: AsyncSession = Depends(get_db)
):
    """Export story as Microsoft Word document"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(story.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle if exists
    if story.subtitle:
        subtitle = doc.add_paragraph(story.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.italic = True
        subtitle_format.size = Pt(14)
    
    # Add story metadata
    doc.add_paragraph()
    metadata = doc.add_paragraph()
    metadata.add_run(f"Genre: ").bold = True
    metadata.add_run(f"{story.genre.value.replace('_', ' ').title()}\n")
    metadata.add_run("Tone: ").bold = True
    metadata.add_run(f"{story.tone.value.title()}\n")
    metadata.add_run("Word Count: ").bold = True
    metadata.add_run(f"{story.word_count:,}")
    
    # Add logline if exists
    if story.logline:
        doc.add_paragraph()
        logline = doc.add_paragraph(story.logline)
        logline_format = logline.runs[0].font
        logline_format.italic = True
    
    # Add page break before chapters
    doc.add_page_break()
    
    # Add each chapter
    for chapter in chapters:
        # Chapter heading
        chapter_heading = doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", level=1)
        
        # Chapter content
        if chapter.content:
            try:
                # Use robust cleaning/parsing
                soup, is_html = clean_and_parse_html(chapter.content)
                
                if not is_html:
                    # Treat as plain text with HTML stripped
                    plain_text = extract_plain_text(chapter.content)
                    paragraphs = plain_text.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text)
                            para_format = para.paragraph_format
                            para_format.first_line_indent = Inches(0.5)
                            para_format.space_after = Pt(12)
                else:
                    # Process HTML content
                    # We iterate over top-level elements
                    for element in iter_html_children(soup):
                        if isinstance(element, NavigableString):
                            text = str(element).strip()
                            if text:
                                doc.add_paragraph(text)
                            continue
                            
                        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                            doc.add_heading(element.get_text(), level=int(element.name[1]))
                            
                        elif element.name == 'p':
                            # Check for images inside paragraph
                            img = element.find('img')
                            if img:
                                src = img.get('src')
                                if src:
                                    try:
                                        image_path = resolve_image_path(src)
                                        if image_path:
                                            doc.add_picture(image_path, width=Inches(6))
                                            # Add caption from alt
                                            caption = img.get('alt')
                                            if caption:
                                                 caption_para = doc.add_paragraph(caption)
                                                 caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                 caption_para.runs[0].font.italic = True
                                        else:
                                            print(f"Image not found at: {image_path} (src: {src})")
                                            doc.add_paragraph(f"[Image missing: {os.path.basename(src)}]")
                                    except Exception as e:
                                        print(f"Error adding image: {e}")
                                        doc.add_paragraph(f"[Image error: {src}]")
                            
                            # Add paragraph text (exclude entire text if it's just an image caption we already handled?? No, simpler to just add text)
                            # But wait, element.get_text() gets ALL text in p.
                            text = element.get_text().strip()
                            # If text exists, add it.
                            if text:
                                para = doc.add_paragraph(text)
                                para.paragraph_format.first_line_indent = Inches(0.5)
                                para.paragraph_format.space_after = Pt(12)

                        elif element.name == 'div':
                             # Similar handling for div containing images
                            img = element.find('img')
                            if img:
                                src = img.get('src')
                                if src:
                                    try:
                                        image_path = resolve_image_path(src)
                                        if image_path:
                                            doc.add_picture(image_path, width=Inches(6))
                                            last_paragraph = doc.paragraphs[-1] 
                                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    except Exception as e:
                                        print(f"Error adding image from div: {e}")
                                            
                        elif element.name == 'img':
                            # Top-level image
                            src = element.get('src')
                            if src:
                                try:
                                    image_path = resolve_image_path(src)
                                    if image_path:
                                        doc.add_picture(image_path, width=Inches(6))
                                        last_paragraph = doc.paragraphs[-1] 
                                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                except Exception as e:
                                    print(f"Error adding image: {e}")

            except Exception as e:
                # Fallback to plain text on error
                print(f"Error processing HTML content: {e}")
                paragraphs = extract_plain_text(chapter.content).split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)
        
        # Add page break between chapters (except for the last one)
        if chapter != chapters[-1]:
            doc.add_page_break()
    
    # Save document to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{story.title.replace(' ', '_')}.docx"
    # Encode filename for non-ASCII characters (RFC 5987)
    encoded_filename = quote(filename, safe='')
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{story_id}/epub")
async def export_epub(
    story_id: UUID,
    db: AsyncSession = Depends(get_db)
):
    """Export story as EPUB ebook"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Create EPUB book
    book = epub.EpubBook()
    
    # Set metadata
    book.set_identifier(str(story.id))
    book.set_title(story.title)
    book.set_language('en')
    
    # Add author (you might want to get this from user)
    book.add_author('NarrativeFlow Author')
    
    # Create title page
    title_content = f'''<html>
    <head>
        <title>{html.escape(story.title)}</title>
    </head>
    <body>
        <h1 style="text-align: center;">{html.escape(story.title)}</h1>
    '''
    
    if story.subtitle:
        title_content += f'<h2 style="text-align: center; font-style: italic;">{html.escape(story.subtitle)}</h2>'
    
    if story.logline:
        title_content += f'<p style="text-align: center; font-style: italic; margin-top: 2em;">{html.escape(story.logline)}</p>'
    
    title_content += f'''<div style="margin-top: 3em;">
        <p><strong>Genre:</strong> {html.escape(story.genre.value.replace('_', ' ').title())}</p>
        <p><strong>Tone:</strong> {html.escape(story.tone.value.title())}</p>
        <p><strong>Word Count:</strong> {story.word_count:,}</p>
    </div>
    </body>
    </html>'''
    
    title_page = epub.EpubHtml(title='Title Page', file_name='title.xhtml', lang='en')
    title_page.content = title_content
    book.add_item(title_page)
    
    # Add chapters
    epub_chapters = []
    image_map: dict = {}
    for chapter in chapters:
        chapter_body = build_epub_body_html(chapter.content or "", book, image_map)
        chapter_content = f'''<html>
        <head>
            <title>Chapter {chapter.number}: {html.escape(chapter.title)}</title>
        </head>
        <body>
            <h1>Chapter {chapter.number}: {html.escape(chapter.title)}</h1>
            {chapter_body}
        </body>
        </html>'''
        
        epub_chapter = epub.EpubHtml(
            title=f"Chapter {chapter.number}: {chapter.title}",
            file_name=f'chapter_{chapter.number}.xhtml',
            lang='en'
        )
        epub_chapter.content = chapter_content
        book.add_item(epub_chapter)
        epub_chapters.append(epub_chapter)
    
    # Define Table of Contents
    book.toc = (
        epub.Link('title.xhtml', 'Title Page', 'title'),
        *[epub.Link(f'chapter_{ch.number}.xhtml', f'Chapter {ch.number}: {ch.title}', f'ch{ch.number}') for ch in chapters]
    )
    
    # Add default NCX and Nav files
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Define CSS style
    style = '''body {
        font-family: Georgia, serif;
        line-height: 1.6;
        margin: 2em;
    }
    h1 {
        text-align: center;
        margin-bottom: 1em;
    }
    p {
        text-indent: 2em;
        margin: 0.5em 0;
    }
    img {
        max-width: 100%;
        height: auto;
        display: block;
        margin: 1em auto;
    }'''
    
    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=style
    )
    book.add_item(nav_css)
    
    # Create spine (reading order)
    book.spine = ['nav', title_page] + epub_chapters
    
    # Write to buffer
    buffer = io.BytesIO()
    epub.write_epub(buffer, book, {})
    buffer.seek(0)
    
    filename = f"{story.title.replace(' ', '_')}.epub"
    encoded_filename = quote(filename, safe='')
    
    return StreamingResponse(
        buffer,
        media_type="application/epub+zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{story_id}/pdf")
async def export_pdf(
    story_id: UUID,
    db: AsyncSession = Depends(get_db)
):
    """Export story as PDF"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Create PDF buffer
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Container for PDF elements
    story_elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='black',
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='grey',
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    chapter_title_style = ParagraphStyle(
        'ChapterTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        spaceBefore=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=12,
        leading=18,
        alignment=TA_JUSTIFY,
        firstLineIndent=0.5*inch,
        spaceAfter=12
    )
    
    # Add title
    story_elements.append(Spacer(1, 0.5*inch))
    story_elements.append(Paragraph(story.title, title_style))
    
    # Add subtitle if exists
    if story.subtitle:
        story_elements.append(Spacer(1, 0.2*inch))
        story_elements.append(Paragraph(story.subtitle, subtitle_style))
    
    # Add metadata
    story_elements.append(Spacer(1, 0.5*inch))
    metadata_text = f"<b>Genre:</b> {story.genre.value.replace('_', ' ').title()}<br/>"
    metadata_text += f"<b>Tone:</b> {story.tone.value.title()}<br/>"
    metadata_text += f"<b>Word Count:</b> {story.word_count:,}"
    story_elements.append(Paragraph(metadata_text, styles['Normal']))
    
    # Add logline if exists
    if story.logline:
        story_elements.append(Spacer(1, 0.3*inch))
        logline_style = ParagraphStyle(
            'Logline',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica-Oblique',
            alignment=TA_CENTER
        )
        story_elements.append(Paragraph(story.logline, logline_style))
    
    # Add page break before chapters
    story_elements.append(PageBreak())
    
    # Add chapters
    for idx, chapter in enumerate(chapters):
        # Chapter title
        chapter_title = f"Chapter {chapter.number}: {chapter.title}"
        story_elements.append(Paragraph(chapter_title, chapter_title_style))
        story_elements.append(Spacer(1, 0.3*inch))
        
        # Chapter content
        if chapter.content:
            soup, is_html = clean_and_parse_html(chapter.content)
            if not is_html:
                paragraphs = extract_plain_text(chapter.content).split('\n')
                for para in paragraphs:
                    if para.strip():
                        safe_para = html.escape(para)
                        story_elements.append(Paragraph(safe_para, body_style))
            else:
                for element in iter_html_children(soup):
                    if isinstance(element, NavigableString):
                        text = str(element).strip()
                        if text:
                            story_elements.append(Paragraph(html.escape(text), body_style))
                        continue

                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        story_elements.append(Paragraph(html.escape(element.get_text()), chapter_title_style))
                        story_elements.append(Spacer(1, 0.2*inch))
                    elif element.name in ['p', 'div']:
                        for img in element.find_all('img'):
                            src = img.get('src')
                            image_path = resolve_image_path(src) if src else None
                            if image_path:
                                img_flowable = build_pdf_image(image_path, 5.5*inch)
                                if img_flowable:
                                    story_elements.append(Spacer(1, 0.1*inch))
                                    story_elements.append(img_flowable)
                                    story_elements.append(Spacer(1, 0.2*inch))
                        text = element.get_text().strip()
                        if text:
                            story_elements.append(Paragraph(html.escape(text), body_style))
                    elif element.name == 'img':
                        src = element.get('src')
                        image_path = resolve_image_path(src) if src else None
                        if image_path:
                            img_flowable = build_pdf_image(image_path, 5.5*inch)
                            if img_flowable:
                                story_elements.append(Spacer(1, 0.1*inch))
                                story_elements.append(img_flowable)
                                story_elements.append(Spacer(1, 0.2*inch))
        
        # Add page break between chapters (except last one)
        if idx < len(chapters) - 1:
            story_elements.append(PageBreak())
    
    # Build PDF
    doc.build(story_elements)
    buffer.seek(0)
    
    filename = f"{story.title.replace(' ', '_')}.pdf"
    encoded_filename = quote(filename, safe='')
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{story_id}/json")
async def export_json(
    story_id: UUID,
    include_metadata: bool = True,
    db: AsyncSession = Depends(get_db)
):
    """Export story as JSON"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    export_data = {
        "title": story.title,
        "subtitle": story.subtitle,
        "logline": story.logline,
        "synopsis": story.synopsis,
        "genre": story.genre.value,
        "tone": story.tone.value,
        "word_count": story.word_count,
        "chapters": [
            {
                "number": ch.number,
                "title": ch.title,
                "content": ch.content,
                "word_count": ch.word_count,
                "summary": ch.summary if include_metadata else None
            }
            for ch in chapters
        ]
    }
    
    if include_metadata:
        # Load characters and plotlines using direct queries
        char_result = await db.execute(
            select(Character).where(Character.story_id == story_id)
        )
        characters = char_result.scalars().all()
        
        plot_result = await db.execute(
            select(Plotline).where(Plotline.story_id == story_id)
        )
        plotlines = plot_result.scalars().all()
        
        export_data["characters"] = [
            {
                "name": char.name,
                "role": char.role.value if char.role else "supporting",
                "personality": char.personality_summary,
                "backstory": char.backstory
            }
            for char in characters
        ]
        
        export_data["plotlines"] = [
            {
                "title": plot.title,
                "description": plot.description,
                "status": plot.status.value if plot.status else "active"
            }
            for plot in plotlines
        ]
    
    return export_data


@router.get("/{story_id}/outline")
async def export_outline(
    story_id: UUID,
    db: AsyncSession = Depends(get_db)
):
    """Export story outline/structure"""
    story = await story_service.get_story(db, story_id)
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    
    chapters = await chapter_service.get_chapters_by_story(db, story_id)
    
    # Load characters and plotlines using direct queries
    char_result = await db.execute(
        select(Character).where(Character.story_id == story_id)
    )
    characters = char_result.scalars().all()
    
    plot_result = await db.execute(
        select(Plotline).where(Plotline.story_id == story_id)
    )
    plotlines = plot_result.scalars().all()
    
    outline = {
        "story": {
            "title": story.title,
            "genre": story.genre.value,
            "tone": story.tone.value,
            "logline": story.logline
        },
        "chapters": [
            {
                "number": ch.number,
                "title": ch.title,
                "summary": ch.summary,
                "outline": ch.outline,
                "key_events": ch.key_events,
                "word_count": ch.word_count,
                "status": ch.status.value
            }
            for ch in chapters
        ],
        "characters": [
            {
                "name": char.name,
                "role": char.role.value if char.role else "supporting",
                "arc": char.arc_description,
                "first_appearance": char.first_appearance_chapter
            }
            for char in characters
        ],
        "plotlines": [
            {
                "title": plot.title,
                "type": plot.type.value if plot.type else "main",
                "status": plot.status.value if plot.status else "active",
                "setup": plot.setup,
                "resolution": plot.resolution
            }
            for plot in plotlines
        ]
    }
    
    return outline
